import os
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_border(cell, **kwargs):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    
    for border_name, border_attrs in kwargs.items():
        border = OxmlElement(f"w:{border_name}")
        for attr_name, attr_value in border_attrs.items():
            border.set(qn(f"w:{attr_name}"), str(attr_value))
        tcPr.append(border)

def set_font(cell, font_name='Calibri (Body)', font_size=14):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), font_name)

# Get the directory of the current script
script_dir = os.path.dirname(os.path.abspath(__file__))

# Construct the path to the image
picture_path = os.path.join(script_dir, 'Images', 'image.png')  # Access the images folder and image.png

# Check if the image path is valid
if not os.path.exists(picture_path):
    print("The specified image path does not exist.")
else:
    # Create a new Document
    doc = Document()
    
    # Add a title to the document
    doc.add_heading('Table Example', level=1)

    # Add a picture to the top right
    doc.add_picture(picture_path, width=Inches(1.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = 2  # Right alignment

    # Create a table with 1 header row and 3 data rows, each with 6 columns
    table = doc.add_table(rows=4, cols=6)

    # Define headers and data
    headers = ["SN", "Model And Describtion", "layout", "Q", "Unit price", "Total price"]

    itemList = []
    quantityList = []
    priceList = []

    with open('requested-data.txt', 'r') as names:
        for line in names:
            itemList.append(line.strip())
    itemList = itemList[1:]
    customerName = itemList[len(itemList)-1]
    itemList = itemList[:-1]

    with open('requested-quantity.txt', 'r') as quantity:
        for line in quantity:
            quantityList.append(line.strip())
    quantityList = quantityList[1:]

    with open('requested-prices.txt', 'r') as prices:
        for line in prices:
            priceList.append(line.strip())
    priceList = priceList[1:]

    data = []
    totalPrice = 0

    for i in range(len(itemList)):
        dataRow = []
        totalPrice += int(quantityList[i]) * int(priceList[i])
        dataRow.extend([i+1, itemList[i], "INSERT PICTURE", quantityList[i], priceList[i], int(quantityList[i]) * int(priceList[i])])
        data.append(dataRow)

    # Populate the header row
    header_row = table.rows[0]
    for col_idx, header in enumerate(headers):
        cell = header_row.cells[col_idx]
        cell.text = header
        
        # Set borders for the header cell
        set_cell_border(
            cell,
            top={"sz": 12, "val": "single", "color": "000000"},
            bottom={"sz": 12, "val": "single", "color": "000000"},
            start={"sz": 12, "val": "single", "color": "000000"},
            end={"sz": 12, "val": "single", "color": "000000"}
        )

        # Set font for the header cell
        set_font(cell, font_name='Calibri (Body)', font_size=14)

    # Populate the data rows
    for row_idx, row_data in enumerate(data, start=1):
        row = table.rows[row_idx]
        for col_idx, cell_data in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = str(cell_data)
            
            # Set borders for the data cell
            set_cell_border(
                cell,
                top={"sz": 12, "val": "single", "color": "000000"},
                bottom={"sz": 12, "val": "single", "color": "000000"},
                start={"sz": 12, "val": "single", "color": "000000"},
                end={"sz": 12, "val": "single", "color": "000000"}
            )

    # Save the document
    doc.save('table_with_header_and_picture.docx')

    print("Document created successfully!")
